"""
Banking RAG DOCX Document Loader.

Reads Word (.docx) banking documents paragraph by paragraph or section by section.
"""

from pathlib import Path
from typing import List, Union

from banking_rag.constants import MetadataKeys
from banking_rag.exceptions import DocumentLoadError
from banking_rag.loaders.base_loader import BaseDocumentLoader, Document
from banking_rag.utils.file_utils import compute_file_hash, get_file_extension, validate_file_exists
from banking_rag.utils.logger import get_logger

logger = get_logger("loaders.docx_loader")


class DocxLoader(BaseDocumentLoader):
    """Loader for Microsoft Word (.docx) banking documents."""

    def supports(self, file_path: Union[str, Path]) -> bool:
        """Checks if file extension is .docx."""
        return get_file_extension(file_path) == ".docx"

    def load(self, file_path: Union[str, Path]) -> List[Document]:
        """Loads a .docx document into Document objects.

        Args:
            file_path: Path to target docx file.

        Returns:
            List of Document objects representing extracted paragraphs/sections.

        Raises:
            DocumentLoadError: If file reading fails.
        """
        path = validate_file_exists(file_path)
        if not self.supports(path):
            raise DocumentLoadError(f"Unsupported file format for DocxLoader: {path}")

        logger.info(f"Loading DOCX document: {path.name}")
        file_hash = compute_file_hash(path)

        try:
            import docx
            doc = docx.Document(str(path))
            full_text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            content = "\n\n".join(full_text)
            
            document = Document(
                content=content,
                file_path=path,
                page_number=1,
                doc_id=f"{file_hash[:12]}_docx",
                metadata={
                    MetadataKeys.DOC_NAME: path.name,
                    MetadataKeys.SOURCE: str(path),
                    MetadataKeys.PAGE_NUMBER: 1,
                    "file_hash": file_hash,
                },
            )
            
            logger.info(f"Successfully loaded DOCX file: {path.name}")
            return [document]

        except ImportError:
            logger.warning("python-docx package not installed.")
            raise DocumentLoadError("DOCX extraction dependency 'python-docx' is not installed.")
        except Exception as e:
            logger.error(f"Error parsing DOCX file {path}: {str(e)}")
            raise DocumentLoadError(f"Failed to extract DOCX content from {path.name}: {str(e)}")
